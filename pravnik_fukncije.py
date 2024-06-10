import os
import requests
import streamlit as st
from bs4 import BeautifulSoup
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.summarize import load_summarize_chain
from langchain_openai.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
import io
import docx
import PyPDF2
from datetime import datetime

from myfunc.varvars_dicts import work_prompts, work_vars

mprompts = work_prompts()

def dl_paragraf(url):
   
    # URL of the webpage you want to scrape se prosledkjuje kao parametar
    # Send a GET request to the webpage
    response = requests.get(url)

    # Check if the request was successful
    if response.status_code == 200:
        # Parse the HTML content of the webpage
        encoding = response.encoding if 'charset' in response.headers.get('content-type', '').lower() else None
        soup = BeautifulSoup(response.content, 'html.parser', from_encoding=encoding)

        # Find all <p> elements with a class that matches "normal" (e.g., "normal", "normal1", "normal2", etc.)
        normal_paragraphs = soup.find_all('p')

        # Create a list to store the extracted text
        text_content = []

        # Iterate through the selected paragraphs, strip HTML tags, and append the text to the list
        for paragraph in normal_paragraphs:
            text = paragraph.get_text(strip=True)
            text_content.append(text)

        # Join the list of text into a single string, separating paragraphs with newlines
        full_text = '\n'.join(text_content)
    else:
        # st.error("Failed to retrieve the web page.")
        
        return full_text

def dl_parlament(url):
     
    response = requests.get(url)

    if response.status_code == 200:
        ime_fajla = url.rsplit('/', 1)[-1]
        if "pdf" in ime_fajla:
            full_text = pdf_from_web(url)
        elif "docx" in ime_fajla:
            full_text = docx_from_web(url)
        else:
            # st.error("Failed to retrieve the web page.")
            full_text = " "
    # else:
        #st.error(f"Failed to download the file {ime_fajla}")

    return full_text
    

# izradjuje sumarizaciju zakona
def sumiraj_zakone(full_text, zakon):
    from langchain.schema import Document
    doc = Document(page_content=full_text)
    
    # st.info("Sumiram zakon: " + zakon)
    # Loading the text document
    #loader = UnstructuredFileLoader(doc, encoding="utf-8")    
    #text_doc = loader.load()
    
    # Initializing ChatOpenAI model
    llm = ChatOpenAI(
          model_name=work_vars["names"]["openai_model"], temperature=0
         )

    chunk_size = 3000
    chunk_overlap = 0
    text_splitter = RecursiveCharacterTextSplitter(
         chunk_size=chunk_size, chunk_overlap=chunk_overlap
      ) 
    
    # pravi dokument od teksta neophodno za rad summarization chaina
    doc = Document(page_content=full_text)
    
    # Splitting the loaded text into smaller chunks
    docs = text_splitter.split_documents([doc])

    # promptovi za prvu i drugu fazu sumarizacije
    prompt_string_pocetak = mprompts["short_summary_begin"]
    
    prompt_string_kraj = mprompts["short_summary_end"]
    
    PROMPT = PromptTemplate(
            template=prompt_string_pocetak, input_variables=["text"]
        )  # Creating a prompt template object
    PROMPT_pam = PromptTemplate(
            template=prompt_string_kraj, input_variables=["text"]
        )  # Creating a prompt template object

    # summarization chain map reduce
    chain = load_summarize_chain(
        llm,
        chain_type="map_reduce",
        verbose=False,
        map_prompt=PROMPT,
        combine_prompt=PROMPT_pam,
        token_max=4000,
    )

    # Run the summarization chain
    sumarizacija = chain.invoke({docs})["output"]
       
    return sumarizacija


# lista zakona od interesa
def lista_zakona():
    search_strings = [
        " O PENZIJSKOM I INVALIDSKOM OSIGURANJU",
        " O ELEKTRONSKOM FAKTURISANJU",
        " O ROKOVIMA IZMIRENJA NOVČANIH OBAVEZA U KOMERCIJALNIM TRANSAKCIJAMA",
        " O POREZU NA DOBIT PRAVNIH LICA",
        " O PRIVREDNIM DRUŠTVIMA",
        " O ZAPOŠLJAVANJU I OSIGURANJU ZA SLUČAJ NEZAPOSLENOSTI",
        " O OBLIGACIONIM ODNOSIMA",
        " O RADU",
        " O CENTRALNOM REGISTRU OBAVEZNOG SOCIJALNOG OSIGURANJA",
        " O JAVNIM NABAVKAMA",
        " O INFORMACIONOJ BEZBEDNOSTI",
        " O MIRNOM REŠAVANJU RADNIH SPOROVA",
        " O UPRAVLJANJU OTPADOM",
        " O ZAPOŠLJAVANJU STRANACA",
        " O CENTRALNOJ EVIDENCIJI STVARNIH VLASNIKA",
        " CARINSKI ",
        " CARINSKOM ",
        " O FISKALIZACIJI",
        " O PORESKOM POSTUPKU I PORESKOJ ADMINISTRACIJI",
        " O PATENTIMA",
        " O POSTUPKU REGISTRACIJE U AGENCIJI ZA PRIVREDNE REGISTRE",
        " O ELEKTRONSKOM DOKUMENTU, ELEKTRONSKOJ IDENTIFIKACIJI I USLUGAMA OD POVERENJA U ELEKTRONSKOM POSLOVANJU",
        " O OSIGURANJU",
        " O SPREČAVANJU PRANJA NOVCA I FINANSIRANJA TERORIZMA",
        " O PARNIČNOM POSTUPKU",
        " O JAVNIM PREDUZEĆIMA",
        " O ELEKTRONSKOJ TRGOVINI",
        " O ZALOŽNOM PRAVU NA POKRETNIM STVARIMA UPISANIM U REGISTAR",
        " O STEČAJU",
        " O INSPEKCIJSKOM NADZORU",
        " O KOLIČINI RASHODA (KALO, RASTUR, KVAR I LOM) NA KOJI SE NE PLAĆA AKCIZA",
        " O OBRAZOVANJU SAVETA ZA MALA I SREDNJA PREDUZEĆA, PREDUZETNIŠTVO I KONKURENTNOST",
                    
    ]

    return search_strings

# skida pdf sa weba i stavlja u string
def pdf_from_web(url):
    response = requests.get(url, stream=True)

    with io.BytesIO(response.content) as f:
        pdf_reader = PyPDF2.PdfReader(f)
        text_stream = io.StringIO()
        for page in pdf_reader.pages:
            text_stream.write(page.extract_text())
    return text_stream.getvalue()

# skida docx sa weba i stavlja u string
def docx_from_web(url):
    response = requests.get(url, stream=True)

    with io.BytesIO(response.content) as f:
        document = docx.Document(f)
        text_stream = io.StringIO()
        for para in document.paragraphs:
            text_stream.write(para.text)
            text_stream.write('\n')
    return text_stream.getvalue()



# pretvara datum string sa sajta u date objekat
def parse_serbian_date(date_string):
    serbian_month_names = {
        "januar": "January",
        "februar": "February",
        "mart": "March",
        "april": "April",
        "maj": "May",
        "jun": "June",
        "jul": "July",
        "avgust": "August",
        "septembar": "September",
        "oktobar": "October",
        "novembar": "November",
        "decembar": "December"
    }

    date_string = date_string.lower()

    for serbian_month, english_month in serbian_month_names.items():
        date_string = date_string.replace(serbian_month, english_month)

    date_string = date_string.strip()

    date_obj = datetime.strptime(date_string, "%d. %B %Y")
    return date_obj